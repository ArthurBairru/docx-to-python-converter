from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import RGBColor
import re

input_file = "input.docx"
doc = Document(input_file)


generated_lines = [
    "from docx import Document",
    "from docx.shared import RGBColor, Pt",
    "from docx.enum.text import WD_ALIGN_PARAGRAPH",
    "from docx.oxml import OxmlElement",
    "from docx.oxml.ns import qn",
    "from docx.enum.style import WD_STYLE_TYPE",
    "",
    "doc = Document()",
    "# Create a character style for default run formatting",
    "default_font_style = doc.styles.add_style('DefaultFont', WD_STYLE_TYPE.CHARACTER)",
    "default_font_style.font.color.rgb = RGBColor(0, 0, 0)",
    "default_font_style.font.size = Pt(8)",
    ""
]

border_functions = [
    "",
    "from docx.oxml.shared import OxmlElement",
    "from docx.oxml.ns import qn",
    "",
    "def set_cell_border(cell, border_name='all', border_size=8, border_color='000000', border_style='single'):",
    "    \"\"\"Set cell borders. border_name: 'top', 'bottom', 'left', 'right', or 'all'\"\"\"",
    "    tcPr = cell._tc.get_or_add_tcPr()",
    "    tcBorders = OxmlElement('w:tcBorders')",
    "    ",
    "    sides = ['top', 'left', 'bottom', 'right'] if border_name == 'all' else [border_name]",
    "    for side in sides:",
    "        border = OxmlElement(f'w:{side}')",
    "        border.set(qn('w:val'), border_style)",
    "        border.set(qn('w:sz'), str(border_size))",
    "        border.set(qn('w:color'), border_color)",
    "        tcBorders.append(border)",
    "    ",
    "    tcPr.append(tcBorders)",
    "    return f\"Cell border {border_name} set: {border_style} {border_size}/8 pt\"",
    "",
    "def set_cell_shading(cell, fill_color='FFFFFF'):",
    "    \"\"\"Set cell background color\"\"\"",
    "    tcPr = cell._tc.get_or_add_tcPr()",
    "    shading = OxmlElement('w:shd')",
    "    shading.set(qn('w:fill'), fill_color)",
    "    shading.set(qn('w:val'), 'clear')",
    "    tcPr.append(shading)",
    "    return f\"Cell shading set: #{fill_color}\"",
    ""
]


# Insert border functions after imports and document creation
generated_lines[12:12] = border_functions

def escape_text(text):
    return text.replace("\\", "\\\\").replace('"', '\\"').replace("\n", "\\n")

def detect_numbering_style_from_text(text):
    """
    Detect numbering or bullet style from paragraph text.
    Supports arabic, letters, roman numerals, bullets.
    """
    if not text.strip():
        return None

    token = text.strip().split()[0]  # first word-like token

    # Arabic numbers: 1. 2) 1.1 1.1.1)
    if re.match(r'^\d+(\.\d+)*[.)]?$', token):
        return "arabic"

    # Letters
    if re.match(r'^[a-z][.)]$', token):
        return "lowercase_letter"
    if re.match(r'^[A-Z][.)]$', token):
        return "uppercase_letter"

    # Roman numerals
    roman_pattern = r'^(?=[MDCLXVI])M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})[.)]?$'
    if re.match(roman_pattern, token, re.IGNORECASE):
        if token.isupper():
            return "roman_uppercase"
        else:
            return "roman_lowercase"

    # Default: bullet
    return "bullet"

def get_appropriate_list_style(numbering_style):
    """
    Map detected numbering style to Word built-in styles
    """
    if numbering_style in ["arabic"]:
        return "List Number"
    elif numbering_style in ["lowercase_letter", "uppercase_letter"]:
        return "List Number 2"
    elif numbering_style in ["roman_lowercase", "roman_uppercase"]:
        return "List Number 3"
    else:
        return "List Bullet"

def get_list_properties(para):
    """
    Returns tuple (is_list, level, numbering_style)
    """
    pPr = para._p.pPr
    if pPr is not None and pPr.numPr is not None:
        ilvl_val = pPr.numPr.ilvl.val if pPr.numPr.ilvl is not None else 0
        numbering_style = detect_numbering_style_from_text(para.text)
        return True, ilvl_val, numbering_style
    return False, 0, None

def remove_numbering_prefix(text):
    """
    Remove common numbering/bullet prefixes
    """
    return re.sub(r'^\s*(\d+(\.\d+)*[.)]?|[A-Za-z][.)]|[IVXLCDM]+[.)])\s+', '', text.strip())

def process_paragraph(para_name, para):
    if not para.text.strip() and len(para.runs) == 0:
        return

    is_list, list_level, numbering_style = get_list_properties(para)
    style = para.style.name if para.style.name != "Normal" else None

    # Add paragraph
    if style:
        generated_lines.append(f"{para_name} = doc.add_paragraph('', style='{style}')")
    else:
        generated_lines.append(f"{para_name} = doc.add_paragraph('')")

    # Apply list style
    if is_list and numbering_style:
        list_style = get_appropriate_list_style(numbering_style)
        generated_lines.append(f"{para_name}.style = '{list_style}'")
        
        # Reduce spacing for list items
        generated_lines.append(f"{para_name}.paragraph_format.space_before = Pt(0)")
        generated_lines.append(f"{para_name}.paragraph_format.space_after = Pt(0)")
        generated_lines.append(f"{para_name}.paragraph_format.line_spacing_rule = 0")  # SINGLE
        generated_lines.append(f"{para_name}.paragraph_format.line_spacing = 1.0")

    # Alignment
    if para.alignment is not None:
        generated_lines.append(f"{para_name}.alignment = WD_ALIGN_PARAGRAPH.{para.alignment.name}")

    # Paragraph formatting
    pf = para.paragraph_format
    if pf.space_before:
        generated_lines.append(f"{para_name}.paragraph_format.space_before = Pt({pf.space_before.pt})")
    if pf.space_after:
        generated_lines.append(f"{para_name}.paragraph_format.space_after = Pt({pf.space_after.pt})")
    if pf.line_spacing:
        if isinstance(pf.line_spacing, float):
            generated_lines.append(f"{para_name}.paragraph_format.line_spacing = {pf.line_spacing}")
        else:
            try:
                generated_lines.append(f"{para_name}.paragraph_format.line_spacing = Pt({pf.line_spacing.pt})")
            except:
                generated_lines.append(f"{para_name}.paragraph_format.line_spacing = 1.0")  # fallback
    else:
        generated_lines.append(f"{para_name}.paragraph_format.line_spacing = 1.0")

     # Detect paragraph bottom border (thick underline)
    try:
        pPr = para._p.pPr
        if pPr is not None and hasattr(pPr, "pBdr") and pPr.pBdr is not None and hasattr(pPr.pBdr, "bottom") and pPr.pBdr.bottom is not None:
            border = pPr.pBdr.bottom
            border_val = getattr(border, "val", "single")
            border_sz = getattr(border, "sz", 4)  # default 0.5 pt
            border_color = getattr(border, "color", "000000")
            generated_lines.append(f"# Paragraph has bottom border (thick underline)")
            generated_lines.append(f"pBorders = OxmlElement('w:pBdr')")
            generated_lines.append(f"bottom = OxmlElement('w:bottom')")
            generated_lines.append(f"bottom.set(qn('w:val'), '{border_val}')")
            generated_lines.append(f"bottom.set(qn('w:sz'), '{border_sz}')")
            generated_lines.append(f"bottom.set(qn('w:color'), '{border_color}')")
            generated_lines.append(f"pBorders.append(bottom)")
            generated_lines.append(f"{para_name}._p.get_or_add_pPr().append(pBorders)")
    except Exception as e:
        generated_lines.append(f"# Error detecting paragraph bottom border: {e}")

    # Process runs
    for run_idx, run in enumerate(para.runs):
        text = escape_text(run.text)
        # Remove numbering prefix only on first run if this is a list
        if is_list and run_idx == 0:
            text = remove_numbering_prefix(text)
        if text:
            generated_lines.append(f"run = {para_name}.add_run({repr(run.text)}, style='DefaultFont')")
            if run.bold:
                generated_lines.append("run.bold = True")
            if run.italic:
                generated_lines.append("run.italic = True")

            if run.font.underline:
                underline_type = run.font.underline
                if isinstance(underline_type, bool):
                    # True/False underline
                    generated_lines.append(f"run.underline = {underline_type}")
                else:
                    # WD_UNDERLINE enum value
                    generated_lines.append(f"run.font.underline = WD_UNDERLINE.{underline_type.name}")

            if run.font.name:
                generated_lines.append(f'run.font.name = "{run.font.name}"')
            if run.font.size:
                size_pt = run.font.size.pt
                generated_lines.append(f"run.font.size = Pt({size_pt})")

def process_table(table_name, table):
    rows = len(table.rows)
    cols = len(table.columns)
    generated_lines.append(f"{table_name} = doc.add_table(rows={rows}, cols={cols})")
    
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            cell_name = f"{table_name}_cell_{i}_{j}"
            generated_lines.append(f"{cell_name} = {table_name}.cell({i}, {j})")
            generated_lines.append(f"{cell_name}.text = ''")
            
            # Detect borders with sizes
            borders = detect_cell_borders(cell)
            if borders.get("has_borders"):
                for side, props in borders.items():
                    if side != "has_borders":
                        # Use the exact size from original table
                        generated_lines.append(
                            f"set_cell_border({cell_name}, border_name='{side}', "
                            f"border_size={props['size']}, border_color='{props['color']}', "
                            f"border_style='{props['style']}')"
                        )
            
            # Process shading if any
            shading = detect_cell_shading(cell)
            if shading:
                generated_lines.append(f'set_cell_shading({cell_name}, "{shading}")')
            
            # Process paragraphs in the cell
            for para_idx, para in enumerate(cell.paragraphs):
                para_var = f"{cell_name}_p{para_idx}"
                process_paragraph(para_var, para)


def detect_cell_borders(cell):
    """Detect cell borders including thickness (sz)."""
    borders_info = {}
    try:
        tcPr = cell._tc.tcPr
        if tcPr is None:
            return {"has_borders": False}
        
        # Convert to XML string and search for borders
        xml_str = tcPr.xml
        if 'w:tcBorders' not in xml_str:
            return {"has_borders": False}
        
        has_visible_borders = False
        
        # Check each side using XML parsing
        from xml.etree import ElementTree as ET
        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        root = ET.fromstring(xml_str)
        
        tcBorders = root.find('.//w:tcBorders', ns)
        if tcBorders is not None:
            for side in ['top', 'left', 'bottom', 'right']:
                border = tcBorders.find(f'.//w:{side}', ns)
                if border is not None:
                    border_style = border.get('{' + ns['w'] + '}val')
                    border_size = border.get('{' + ns['w'] + '}sz')
                    border_color = border.get('{' + ns['w'] + '}color')
                    
                    if border_size is not None:
                        border_size = int(border_size)
                    else:
                        border_size = 8
                    
                    if border_color is None:
                        border_color = '000000'
                    
                    if (border_style not in [None, 'nil', 'none'] and 
                        border_size is not None and border_size > 0):
                        has_visible_borders = True
                        borders_info[side] = {
                            'style': border_style,
                            'size': border_size,
                            'color': border_color
                        }
        
        borders_info["has_borders"] = has_visible_borders
        return borders_info
        
    except Exception as e:
        print(f"Error detecting borders: {e}")  # Debug output
        return {"has_borders": False, "error": str(e)}


def set_cell_border(cell, border_name='all', border_size=8, border_color='000000', border_style='single'):
    """
    Set cell borders. border_name: 'top', 'bottom', 'left', 'right', or 'all'
    border_size: size in eighths of a point (1 = 1/8 pt)
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    sides = ['top', 'left', 'bottom', 'right'] if border_name == 'all' else [border_name]
    for side in sides:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), border_style)
        border.set(qn('w:sz'), str(border_size))
        border.set(qn('w:color'), border_color)
        tcBorders.append(border)

    tcPr.append(tcBorders)
    return f"# Cell border set: {border_name} {border_style} {border_size}/8 pt"



def set_cell_shading(cell, fill_color="FFFFFF"):
    """Set cell background color."""
    tcPr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), fill_color)
    shading.set(qn('w:val'), 'clear')
    tcPr.append(shading)
    return f"# Cell shading set: #{fill_color}"

def detect_cell_shading(cell):
    """Detect cell background color."""
    try:
        tcPr = cell._tc.tcPr
        if tcPr is None:
            return None
            
        shading = tcPr.shd
        if shading is not None:
            fill_color = getattr(shading, 'fill', None)
            if fill_color and fill_color != 'auto' and fill_color != 'FFFFFF':
                return fill_color
        return None
    except:
        return None

# Iterate over body elements
para_counter = 0
table_counter = 0


for child in doc.element.body:
    if isinstance(child, CT_Tbl):
        table_counter += 1
        table = doc.tables[table_counter - 1]
        process_table(f"table{table_counter}", table)
    elif isinstance(child, CT_P):
        para_counter += 1
        para = doc.paragraphs[para_counter - 1]
        process_paragraph(f"p{para_counter}", para)

generated_lines.append('doc.save("output.docx")')

with open("recreate_docx.py", "w", encoding="utf-8") as f:
    f.write("\n".join(generated_lines))

print("Python code generated in recreate_docx.py")
